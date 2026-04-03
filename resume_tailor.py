"""
AI Resume Tailoring Engine.

Generates custom-tailored resumes for each job using the JD + master CV.
Creates PDF files that can be uploaded automatically during Easy Apply.
"""

import logging
import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Optional

log = logging.getLogger("lla.resume_tailor")


class ResumeTailor:
    """Generate tailored resumes per job using AI."""

    def __init__(self, ai, cfg: dict):
        self.ai = ai
        self.cfg = cfg
        rt_cfg = cfg.get("resume_tailoring", {})
        self.enabled = rt_cfg.get("enabled", False)
        self.master_resume_path = rt_cfg.get("master_resume_path", "")
        self.master_resume_text = rt_cfg.get("master_resume_text", "")
        self.output_dir = rt_cfg.get("output_dir", "data/tailored_resumes")
        self.output_format = rt_cfg.get("format", "pdf")
        self.template_style = rt_cfg.get("template_style", "professional")

        if self.enabled:
            Path(self.output_dir).mkdir(parents=True, exist_ok=True)
            self._load_master_resume()

    def _load_master_resume(self):
        """Load master resume text from file or config."""
        if self.master_resume_text:
            return

        # Try from AI config cv_text
        cv_text = self.cfg.get("ai", {}).get("cv_text", "")
        if cv_text:
            self.master_resume_text = cv_text
            return

        cv_file = self.cfg.get("ai", {}).get("cv_text_file", "")
        if cv_file and Path(cv_file).exists():
            try:
                self.master_resume_text = Path(cv_file).read_text(encoding="utf-8")
                return
            except Exception as e:
                log.warning(f"Could not read CV file: {e}")

        if self.master_resume_path and Path(self.master_resume_path).exists():
            try:
                # Try reading as text (for .txt files)
                if self.master_resume_path.endswith('.txt'):
                    self.master_resume_text = Path(self.master_resume_path).read_text(encoding="utf-8")
                else:
                    log.info(f"Master resume is binary ({self.master_resume_path}), using AI cv_text for tailoring")
            except Exception as e:
                log.warning(f"Could not read master resume: {e}")

        if not self.master_resume_text:
            # Fall back to profile context from AI
            if self.ai and hasattr(self.ai, 'profile_context'):
                self.master_resume_text = self.ai.profile_context
                log.info("Using AI profile context as master resume text")

    def tailor_resume(self, job_title: str, company: str, description: str,
                      match_result: dict = None) -> Optional[str]:
        """
        Generate a tailored resume for a specific job.

        Returns:
            File path to the generated resume, or None if generation fails.
        """
        if not self.enabled or not self.ai or not self.ai.enabled:
            return None

        if not self.master_resume_text:
            log.warning("No master resume text available for tailoring")
            return None

        log.info(f"   📝 Tailoring resume for {job_title} @ {company}...")

        # Generate tailored content using AI
        tailored_content = self._generate_tailored_content(
            job_title, company, description, match_result
        )
        if not tailored_content:
            return None

        # Generate file
        safe_company = re.sub(r'[^\w\s-]', '', company)[:30].strip()
        safe_title = re.sub(r'[^\w\s-]', '', job_title)[:30].strip()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_company}_{safe_title}_{timestamp}"

        if self.output_format == "pdf":
            return self._generate_pdf(tailored_content, filename)
        elif self.output_format == "docx":
            return self._generate_docx(tailored_content, filename)
        else:
            return self._generate_txt(tailored_content, filename)

    def _generate_tailored_content(self, job_title: str, company: str,
                                   description: str, match_result: dict = None) -> Optional[dict]:
        """Use AI to generate tailored resume sections."""
        skill_info = ""
        if match_result:
            matches = match_result.get("skill_matches", [])
            missing = match_result.get("missing_skills", [])
            if matches:
                skill_info += f"\nMatching skills to emphasize: {', '.join(matches)}"
            if missing:
                skill_info += f"\nSkills to address if possible: {', '.join(missing)}"

        system_prompt = f"""You are a professional resume writer. Rewrite the candidate's resume to be tailored for a specific job.

RULES:
- Keep all factual information accurate — do NOT fabricate experience, skills, or achievements
- Reorder and emphasize skills/experience that match the job requirements
- Adjust the professional summary to target this specific role
- Use keywords from the job description naturally
- Keep the same structure: Summary, Experience, Education, Skills, Certifications
- Be concise — fit on 1-2 pages maximum
- Output in plain text with clear section headers (ALL CAPS)

MASTER RESUME:
{self.master_resume_text[:4000]}
{skill_info}"""

        user_prompt = f"""Tailor this resume for:
Job Title: {job_title}
Company: {company}
Job Description: {description[:2000]}

Output the complete tailored resume in plain text format with section headers."""

        try:
            old_max = self.ai.max_tokens
            self.ai.max_tokens = 2000
            result = self.ai._call_llm(system_prompt, user_prompt)
            self.ai.max_tokens = old_max

            if result and len(result) > 100:
                return self._parse_resume_sections(result)
        except Exception as e:
            log.warning(f"Resume tailoring AI call failed: {e}")

        return None

    def _parse_resume_sections(self, text: str) -> dict:
        """Parse AI output into resume sections."""
        sections = {
            "full_text": text,
            "name": "",
            "summary": "",
            "experience": "",
            "education": "",
            "skills": "",
            "certifications": "",
        }

        # Extract name from personal config
        personal = self.cfg.get("personal", {})
        sections["name"] = personal.get("full_name", "")

        # Split by section headers
        current = "summary"
        lines = text.split("\n")
        section_map = {
            "summary": ["summary", "professional summary", "profile", "objective"],
            "experience": ["experience", "work experience", "professional experience", "employment"],
            "education": ["education", "academic", "qualifications"],
            "skills": ["skills", "technical skills", "core competencies", "technologies"],
            "certifications": ["certifications", "certificates", "professional development"],
        }

        for line in lines:
            line_lower = line.strip().lower()
            matched = False
            for section_key, keywords in section_map.items():
                if any(kw == line_lower or line_lower.startswith(kw) for kw in keywords):
                    current = section_key
                    matched = True
                    break
            if not matched:
                sections[current] = sections.get(current, "") + line + "\n"

        return sections

    def _generate_pdf(self, content: dict, filename: str) -> Optional[str]:
        """Generate a PDF resume using fpdf2."""
        try:
            from fpdf import FPDF
        except ImportError:
            log.warning("fpdf2 not installed. Run: pip install fpdf2. Falling back to text.")
            return self._generate_txt(content, filename)

        filepath = os.path.join(self.output_dir, f"{filename}.pdf")

        try:
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=20)
            pdf.add_page()

            # Name header
            name = content.get("name", "")
            if name:
                pdf.set_font("Helvetica", "B", 18)
                pdf.cell(0, 12, name, ln=True, align="C")
                pdf.ln(2)

            # Contact info
            personal = self.cfg.get("personal", {})
            contact_parts = []
            if personal.get("email"):
                contact_parts.append(personal["email"])
            if personal.get("phone"):
                contact_parts.append(personal["phone"])
            if personal.get("city"):
                loc = personal["city"]
                if personal.get("country"):
                    loc += f", {personal['country']}"
                contact_parts.append(loc)
            if contact_parts:
                pdf.set_font("Helvetica", "", 9)
                pdf.cell(0, 6, " | ".join(contact_parts), ln=True, align="C")
                pdf.ln(4)

            # Sections
            full_text = content.get("full_text", "")
            if full_text:
                self._write_resume_text(pdf, full_text)
            else:
                for section_key in ["summary", "experience", "education", "skills", "certifications"]:
                    text = content.get(section_key, "").strip()
                    if text:
                        pdf.set_font("Helvetica", "B", 12)
                        pdf.cell(0, 8, section_key.upper(), ln=True)
                        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
                        pdf.ln(2)
                        pdf.set_font("Helvetica", "", 10)
                        pdf.multi_cell(0, 5, text)
                        pdf.ln(3)

            pdf.output(filepath)
            log.info(f"   📄 Tailored resume saved: {filepath}")
            return filepath

        except Exception as e:
            log.warning(f"PDF generation failed: {e}")
            return self._generate_txt(content, filename)

    def _write_resume_text(self, pdf, text: str):
        """Write resume text to PDF with basic formatting."""
        section_headers = ["summary", "professional summary", "experience",
                          "work experience", "education", "skills",
                          "technical skills", "certifications", "certificates",
                          "core competencies", "professional experience"]

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                pdf.ln(2)
                continue

            # Check if it's a section header
            if stripped.upper() == stripped and len(stripped) > 3 and any(
                h in stripped.lower() for h in section_headers
            ):
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 8, stripped, ln=True)
                pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
                pdf.ln(2)
            elif stripped.startswith("- ") or stripped.startswith("• "):
                pdf.set_font("Helvetica", "", 10)
                pdf.cell(5)
                pdf.multi_cell(0, 5, stripped)
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.multi_cell(0, 5, stripped)

    def _generate_docx(self, content: dict, filename: str) -> Optional[str]:
        """Generate a DOCX resume using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            log.warning("python-docx not installed. Run: pip install python-docx. Falling back to text.")
            return self._generate_txt(content, filename)

        filepath = os.path.join(self.output_dir, f"{filename}.docx")

        try:
            doc = Document()

            # Name
            name = content.get("name", "")
            if name:
                p = doc.add_heading(name, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact
            personal = self.cfg.get("personal", {})
            contact_parts = []
            if personal.get("email"):
                contact_parts.append(personal["email"])
            if personal.get("phone"):
                contact_parts.append(personal["phone"])
            if contact_parts:
                p = doc.add_paragraph(" | ".join(contact_parts))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Full text
            full_text = content.get("full_text", "")
            if full_text:
                for line in full_text.split("\n"):
                    stripped = line.strip()
                    if not stripped:
                        continue
                    if stripped.upper() == stripped and len(stripped) > 3:
                        doc.add_heading(stripped.title(), level=1)
                    elif stripped.startswith("- ") or stripped.startswith("• "):
                        doc.add_paragraph(stripped[2:], style="List Bullet")
                    else:
                        doc.add_paragraph(stripped)

            doc.save(filepath)
            log.info(f"   📄 Tailored resume saved: {filepath}")
            return filepath

        except Exception as e:
            log.warning(f"DOCX generation failed: {e}")
            return self._generate_txt(content, filename)

    def _generate_txt(self, content: dict, filename: str) -> Optional[str]:
        """Fallback: generate a plain text resume."""
        filepath = os.path.join(self.output_dir, f"{filename}.txt")
        try:
            text = content.get("full_text", "")
            if not text:
                # Build from sections
                parts = []
                name = content.get("name", "")
                if name:
                    parts.append(name.upper())
                    parts.append("=" * len(name))
                for key in ["summary", "experience", "education", "skills", "certifications"]:
                    sec = content.get(key, "").strip()
                    if sec:
                        parts.append(f"\n{key.upper()}\n{'-' * len(key)}")
                        parts.append(sec)
                text = "\n".join(parts)

            Path(filepath).write_text(text, encoding="utf-8")
            log.info(f"   📄 Tailored resume saved: {filepath}")
            return filepath
        except Exception as e:
            log.warning(f"Text resume generation failed: {e}")
            return None
